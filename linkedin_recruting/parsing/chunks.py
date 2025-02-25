import os
import sys
from docx2pdf import convert
import re
# from pathlib import Path
from dotenv import load_dotenv
# from langchain_community.vectorstores import Chroma
# from langchain_community.vectorstores import Qdrant

# from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter

from langchain_community.document_loaders import PyMuPDFLoader

from PyPDF2 import PdfReader

from helper import generate_random_date, generate_random_id
from helper import isFreelance, extractEmail, anonymizeResume
from datetime import datetime

from libs import setLLM

from langchain_community.vectorstores import Qdrant
from langchain_openai import OpenAIEmbeddings
import fitz

from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams




# Add the subfolders to the Python path
sys.path.append(os.path.join(os.path.dirname(__file__), '../'))
sys.path.append(os.path.join(os.path.dirname(__file__), '../llm'))
sys.path.append(os.path.join(os.path.dirname(__file__), '../email'))
sys.path.append(os.path.join(os.path.dirname(__file__), '../mysqldb'))

from mails import sendEmailGeneral
from prompts import extractExperienceCandidat, extractDiplomeCandidat, extractHardSkillsCandidat, extractCertificationsCandidat, extractPhoneCandidat
from prompts import extractExperienceRequired, extractDiplomeRequired, extractHardSkillsRequired, extractCertificationsRequired
from prompts import extractDiplomeAnneePhone
import logging

from libs import Application, Job, Task
from typing import List

load_dotenv()
resume_folder = os.environ['RESUME_FOLDER']
llm_type = os.environ['LLM_TYPE']

# We create resume_folder if it does not exist
if not os.path.exists(resume_folder):
    os.makedirs(resume_folder)

def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text



def processSingleJob(job_pdf_path: str, task: str, llm) -> Job:

    if job_pdf_path == None:
        logging.info("Invalid input file. The can't be None")
        raise Exception("Invalid input file. The can't be None")

    elif not os.path.exists(job_pdf_path):
        logging.info(f"Invalid path. Can not find {job_pdf_path}")
        raise Exception(f"Invalid path. Can not find {job_pdf_path}")  

    job = Job(job_pdf_path=job_pdf_path, date=generate_random_date(), taskId=task.Id)

    logging.info(f"Job initalized with task = {job.taskId}")


    # Checking if the role is already in the database
    #If the role is already in the database we skip it

    if job.roleId == None:
        logging.info(f"...............=>Job {job.role} is not in the database yet. We generate roleId")
        job.roleId = setRoleId(job.role)

    else:
        logging.info(f"...............=>Job {job.role} is already in the database with roleID = {job.roleId}")
        return job, True

        
    # Extract chunk 
    job_desc_data= getChunk(job_pdf_path)[0].page_content
    logging.info(f"...............=>Job {job.role} resume data read")

    logging.info("")
    logging.info(f"...............=>Job {job.role} Extracting requirements")
    logging.info("")

    job.experience = extractExperienceRequired(context=job_desc_data, llm=llm)
    logging.info("")
    logging.info(f"...............=>Job {job.role} experience = {job.experience}")
    logging.info("")

    job.diplome = extractDiplomeRequired(context=job_desc_data, llm=llm)
    logging.info("")
    logging.info(f"...............=>Job {job.role} diplome = {job.diplome}")
    logging.info("")

    job.certifications =extractCertificationsRequired(context=job_desc_data, llm=llm)
    logging.info("")
    logging.info(f"...............=>Job {job.role} certifications = {job.certifications}")
    logging.info("")

    job.hard_skills = extractHardSkillsRequired(context=job_desc_data, llm=llm)
    logging.info("")
    logging.info(f"...............=>Job {job.role} Hard Skills = {job.hard_skills}")
    logging.info("")

    job.soft_skills = "NONE"
    job.langues = "NONE"

    # Saving job to mysql database
    logging.info(f"...............=>Saving job {job.role}")
    job.save()
    logging.info(f"...............=>Saved job {job.role}")

    return job, False

 
    
def getChunk(file_path):

    '''
    Generate single chunk from a pdf document

    Input: 
        - file_path: input PDF file 

    Output: 
        - chunks: 


    '''    
    
    # loader = PyPDFLoader(file_path)
    # documents = loader.load()

    loader = PyMuPDFLoader(file_path)
    documents = loader.load()

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=50000, chunk_overlap=0)  # Set high chunk size to avoid splitting

    chunks = text_splitter.split_documents(documents)

    if len(chunks) ==0 or len(documents)==0:

        logging.info(f"Issue when spliting file {file_path}. We have only {len(chunks)} chunks and {len(documents)} documents")

        raise Exception(f"Issue when spliting file {file_path}. We have only {len(chunks)} chunks and {len(documents)} documents")
        
    logging.info(f"Number of chunks={len(chunks)}")

    # Putting all content in one chunk
    if len(chunks) > 1:
        for i in range(1, len(chunks)):
            chunks[0].page_content = chunks[0].page_content + chunks[i].page_content

   
        
    # Deleting other chunks
    chunks = [chunks[0]]
    

    
    return chunks



def setRoleId(role):

    '''
        - This function is used to generate RoleID from Role name.
        - Later we will use random generation

    '''
    
    role_id = {"Tech Lead Data Engineering": "3kyBidhu",
            "Generative AI Engineer": "7hZ6glHq",
            "Consultant Data Power BI": "AIegN6My",
            "Consultant Teradata": "hcve6Ik8",
            "Consultant Data Management": "I0tLuRGw",
            "ML Engineer": "ieU9peuC",
            "Lead Power BI": "mIYDb0JL",
            "Data Engineer": "NBOqKK9H",
            "Consultant Data Qlick": "NJvt5YQB",
            "Consultant Data Qlik": "AS5fl5v0",
            "Consultant Data Integration – Informatica _ Talend": "pEumyRVP",
            "Consultant ETL IBM DataStage": "s4YIUtdV",
            "Data Solutions Architect": "KgY8qMd4",
            "Analytics Engineer": "0q7xgtZl",
            "Product Owner Data": "qwerty0"}


    if role in role_id:
        logging.info(f"role {role} already have a roleId which is {role_id[role]}")
        return role_id[role]

    else:
        logging.info(f"{role} does not have a roleId yet. We generate a new one")
        roleId = generate_random_id
        logging.info(f"The generated roleId is {roleId}")
        return roleId


def processSingleApplication(msg_file_path, task, llm):

    logging.info("Inside function processSingleApplication")
    logging.info(f"msg_file_path={msg_file_path}")


    '''
    Process a job application.

    Input: 
        - msg_file_path: msg file representing the CV of the candidate

    Output: 
        - applicationData: Json reprensenting the qualification of the candidate


    '''
    try:
        # Getting application metadata

        if not os.path.exists(msg_file_path):
            logging.error(f"File {msg_file_path} does not exist. Function processApplication")
            raise Exception(f"File {msg_file_path} does not exist")
        
        application = Application(msg_file_path)
        logging.info("")
        logging.info(f"Name = {application.name}")
        logging.info(f"Position = {application.role}")

        # We check if the candidate is already in the database
        if application.isApplicationOld:
            logging.info("")
            logging.info("")
            logging.info(f"Candidate {application.name} has already applied to {application.role} position.")
            logging.info("")
            return application 
  

        application.isApplicationOld = False

        logging.info("")
        logging.info("Reading downloaded resume")
        # We extract text from resume
        laparams = LAParams(line_overlap=0.5, detect_vertical=True, all_texts=True)

        # We get the extension of resume
        _, extension = os.path.splitext(application.pathResume)
        

        if extension == ".pdf":
            text = extract_text(application.pathResume, laparams=laparams)
        elif extension == ".docx" or extension == ".doc":
            logging.info("-----------------------------")
            logging.info(f"Attachement extension is word {extension}")
            logging.info("-----------------------------")
            text = extract_text_from_docx(application.pathResume)
        else:
            raise Exception("Invalid file extension when extracting text from resume in processSingleApplication")

        text = nettoyer_texte(text)
        logging.info("")
        logging.info("Content extracted from resume")
        logging.info("")


        # We check the length of extracted text
        threshold = 100
        if len(text) <=threshold and extension==".pdf":

            logging.info("")

            logging.info(f"Did not extract text because len(text) < {threshold}")

            logging.info("Trying another method with package PdfReader")

            text = extract_text_from_pdf(application.pathResume)
            
            if len(text) <= threshold:
                
        
                logging.info("")
                logging.info("")

                raise Exception(f"Still Unable to extract data from the resume of {application.name}. len(text) <={threshold}. Check the pdf resume of the applicant")


        # Getting email, freelance status
        email = extractEmail(text)

        if len(email) >= 1:
            application.email = email[0]
        else:
            logging.info("")
            logging.info("No email found")
            application.email = None

        if isFreelance(text):
            application.freelance = "Yes"
            logging.info("")
            logging.info("We got a freelance !")
        else:
            application.freelance = "NO"
        logging.info("")
        logging.info(f"email={application.email} and freelance={application.freelance}")

        anonymised_resume = anonymizeResume(fullname=application.name, resume=text)

        # Extracting degree, year and phone number
        logging.info("")
        logging.info("Extracting diplome, annee and numero ...")
        diplome, annee_diplome, phone = extractDiplomeAnneePhone(anonymised_resume, llm=llm)
        logging.info(f"diplome={diplome}  annee={annee_diplome}  and phone={phone}")

        application.phone = phone
        application.diplome = diplome
        application.annee_diplome = annee_diplome



        # Extracting experience
        logging.info("")
        logging.info("Extracting experience ...")
        application.experience = extractExperienceCandidat(anonymised_resume, llm=llm, annee=application.annee_diplome)
        logging.info(f"Candidate diplome={application.diplome} Candidate graduation={application.annee_diplome} Candidate experience={application.experience} .")

         # Extracting experience
        logging.info("")
        logging.info("Extracting hard skills ...")
        application.hard_skills = extractHardSkillsCandidat(anonymised_resume, llm=llm)

    

        # application.certifications = extractCertificationsCandidat(text, llm=llm)
        # logging.info(f"Candidate certifications={application.certifications} .")
        application.certifications = "NONE"

        # Loading job data
        logging.info("")
    
        # logging.info(f"Loading corresponding an empty job")
        job = Job(taskId=task.Id)
        roleId=application.roleId
        # logging.info(f"Loading corresponding job for roleId={roleId}")
        job.load(roleId=roleId)
        # logging.info(f"Loaded job role Id={job.roleId} Job experience={job.experience} Degree = {job.diplome}")

        # Calculatin application score
        application.scoring(job.experience, job.diplome)

        logging.info("")
        logging.info(f"Required exp={job.experience} Required Degree = {job.diplome}")
        logging.info(f"{application.score} score= {application.score} exp= {application.experience }  degree={application.diplome}")

        # Extract qualifications
        
        #**********************************************************
        #**********************************************************
        #**********************************************************
        #                  Abdelaziz Jaddi ICI
        #  def function_abdelaziz(context):
        #   - Input: contex[str]: Data of the resume
        #   - Output: experience[int]: experience
        #
        #   return experience
        #**********************************************************
        #**********************************************************
        # applicationData["experience"] = function_abdelaziz(context=chunks_application[0].page_content)

        # application.certifications = "None"
        # application.hard_skills = "None"
        application.soft_skills = "None"
        application.langues = "None"


        # Scoring
        # application.score = 100
        application.alternative_score = "None"
        application.alternative_role = "None"
        #**************************************************************
        #**************************************************************
        #**************************************************************
        application.save()
        logging.info(f"Saved application of {application.name} for role {application.role}.")


        return application

    except Exception as e:

        logging.info("")
        logging.info(f"In function processSingleApplication, issue with job application {msg_file_path}")
        logging.info(f"Error = {e}")

        raise Exception(e)



def nettoyer_texte(texte):
    # Enlever les retours à la ligne inutiles (les conserver uniquement entre les sections)
    texte = re.sub(r'\n+', '\n', texte)
    
    # Enlever les espaces au début et à la fin
    texte = texte.strip()

    return texte


def extract_text_from_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    text = []
 
    # # Extraire le texte des paragraphes normaux
    for para in doc.paragraphs:
        text.append(para.text)
 
    # Extraire le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append("\t".join(row_text))  # Format tabulé pour mieux séparer les colonnes
 
    return "\n".join(text)

