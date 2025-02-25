from docx import Document
 
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
 
    # # Extraire le texte des paragraphes normaux
    for para in doc.paragraphs:
        text.append(para.text)
 
    # Extraire le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append("\t".join(row_text))  # Format tabulé pour mieux séparer les colonnes
 
    return "\n".join(text)
 
# Exemple d'utilisation
file_path = "C://Users\gaelk\OneDrive\Desktop\Jobs\July 2023/Aubay/CV_KAMDEM_Gael_BA.docx"
texte = extract_text_from_docx(file_path)
print(texte)